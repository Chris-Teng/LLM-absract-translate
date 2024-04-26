import PyPDF2,requests,os,glob,argparse,time,email,re
from openpyxl import Workbook
from docx import Document
from requests.adapters import HTTPAdapter, Retry
from bs4 import BeautifulSoup

def readPdfFile(filename):
    reader = PyPDF2.PdfReader(filename)
    num_pages = len(reader.pages)
    count = 0
    text = ""
    while count < num_pages:
        page = reader.pages[count]
        count += 1
        text += page.extract_text()
    return text
def readWordFile(filename):
    """Read contents of a Microsoft Word (.doc/.docx) file."""
    try:
        # document = Document(os.path.join("files", filename))   # Open the document
        document = Document(filename)                # Open the document
        fullText = []                                # Create empty list to store lines
        
        for para in document.paragraphs:             # Iterate through paragraphs
            fullText.append(para.text)               # Append each line to the list
    
        return '\n'.join(fullText)                   # Return the joined string
    except Exception as e:
        print(f"Error reading {filename}: {e}")      # Handle errors
        return None
def readTxtFile(filename):
    with open(filename, "r") as f:
        text = f.read()
    return text
def readEmlFile(filename):
  """
  读取 eml 文件并获取邮件内容

  Args:
    filename: eml 文件名

  Returns:
    邮件内容
  """
  with open(filename, 'rb') as f:
    # 解析 eml 文件
    msg = email.message_from_bytes(f.read())

    # 获取邮件正文
    if msg.is_multipart():
      # 处理多段邮件
        body = ''
        for part in msg.walk():
            if part.get_content_maintype() == 'text':
            # 去除 HTML 标签
                soup = BeautifulSoup(part.get_payload(decode=True), 'html.parser')
                text = soup.get_text(separator='\n')
                body += text
                body = body.replace('\n', ' ')
    else:
      # 处理单段邮件
        body = re.sub(r'<[^>]*>', '', msg.get_payload(decode=True).decode())
    # 返回邮件内容
    return body
if __name__ == "__main__":
    wb = Workbook()
    ws = wb.active

    # Setup Arguments Parser
    ap = argparse.ArgumentParser()  
    ap.add_argument("directory", nargs="?", default=None, help="Directory Path Containing Documents")
    args = vars(ap.parse_args())

    # Validate Input Parameters
    if not args["directory"]:
        print("\nUsage:\npython main.py /path/to/input/dir/\n")
        exit()

    dirName = args['directory']
    if not os.path.exists(dirName):
        print("Invalid Directory Path Specified!\n")
        exit()

    # Extract All Supported Formats
    supportedFormats = ['*.pdf', '*.docx', '*.doc' , '*.txt', '*.eml']
    foundFilesList = []

    for extFormat in supportedFormats:
       foundFilesList.extend([foundFile for foundFile in glob.glob(dirName+'/**/'+extFormat, recursive=True)])

    # Process Found Files Accordingly
    for foundFile in sorted(set(foundFilesList)):
        _, fileExtension = os.path.splitext(foundFile)

        if fileExtension == '.pdf':
            processedContent = readPdfFile(foundFile)
        elif fileExtension == '.docx' or fileExtension == '.doc':
            processedContent = readWordFile(foundFile)
        elif fileExtension == '.txt':
            processedContent = readTxtFile(foundFile)
        elif fileExtension == '.eml':
            processedContent = readEmlFile(foundFile)

        # print('\nProcessed Content from "%s"\n%s' % (foundFile, processedContent))
        print('\nProcessed Content from "%s"\n' % (foundFile))

        if processedContent:
            
            s = requests.Session()
            retries = Retry(total=10,
                            backoff_factor=0.1,
                            status_forcelist=[ 500, 502, 503, 504 ])

            s.mount('http://', HTTPAdapter(max_retries=retries))
            url = "http://127.0.0.1:5000/v1/completions"

            headers = {
                "Content-Type": "application/json"
            }
            data = {
                "prompt": f"Fully summarize the text delimited by triple backticks.```{processedContent[:10000]}```",
                "max_tokens": 4096,   # 一个token约等于4个单词
                "temperature": 1,
                "top_p": 0.9,
                "seed": 10
            }
            try:
                response = s.post(url, headers=headers, json=data, timeout=3600)
                absText = response.json()['choices'][0]['text']
            except:
                print("请求失败，请重启模型")
                input("按任意键继续")
                try:
                    response = s.post(url, headers=headers, json=data, timeout=600)
                    absText = response.json()['choices'][0]['text']
                except:
                    ws.append([foundFile, '运行时间过长，停止运行'])
                    wb.save('output.xlsx')
                    continue
            try:
                ws.append([foundFile, absText])
            except:
                ws.append([foundFile, '结果写入有误，文件内容可能无法读取，请检查文件内容'])
        wb.save('output.xlsx')