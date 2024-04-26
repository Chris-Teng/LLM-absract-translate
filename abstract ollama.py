import PyPDF2,requests,os,glob,argparse
from openpyxl import Workbook
from docx import Document
def readPdfFile(filename):
    reader = PyPDF2.PdfReader(filename)
    num_pages = len(reader.pages)
    count = 0
    text = ""
    while count < num_pages:
        page = reader.pages[count]
        count += 1
        text += page.extract_text()
    return text
def readWordFile(filename):
    """Read contents of a Microsoft Word (.doc/.docx) file."""
    try:
        # document = Document(os.path.join("files", filename))   # Open the document
        document = Document(filename)                # Open the document
        fullText = []                                # Create empty list to store lines
        
        for para in document.paragraphs:             # Iterate through paragraphs
            fullText.append(para.text)               # Append each line to the list
    
        return '\n'.join(fullText)                   # Return the joined string
    except Exception as e:
        print(f"Error reading {filename}: {e}")      # Handle errors
        return None
def readTxtFile(filename):
    with open(filename, "r") as f:
        text = f.read()
    return text

if __name__ == "__main__":
    wb = Workbook()
    ws = wb.active


    # Setup Arguments Parser
    ap = argparse.ArgumentParser()
    ap.add_argument("directory", nargs="?", default=None, help="Directory Path Containing Documents")
    args = vars(ap.parse_args())

    # Validate Input Parameters
    if not args["directory"]:
        print("\nUsage:\npython main.py /path/to/input/dir/\n")
        exit()

    dirName = args['directory']
    if not os.path.exists(dirName):
        print("Invalid Directory Path Specified!\n")
        exit()

    # Extract All Supported Formats
    supportedFormats = ['*.pdf', '*.docx', '*.doc' , '*.txt']
    foundFilesList = []

    for extFormat in supportedFormats:
       foundFilesList.extend([foundFile for foundFile in glob.glob(dirName+'/**/'+extFormat, recursive=True)])

    # Process Found Files Accordingly
    for foundFile in sorted(set(foundFilesList)):
        _, fileExtension = os.path.splitext(foundFile)

        if fileExtension == '.pdf':
            processedContent = readPdfFile(foundFile)
        elif fileExtension == '.docx' or fileExtension == '.doc':
            processedContent = readWordFile(foundFile)
        elif fileExtension == '.txt':
            processedContent = readTxtFile(foundFile)

        # print('\nProcessed Content from "%s"\n%s' % (foundFile, processedContent))
        print('\nProcessed Content from "%s"\n' % (foundFile))

        if processedContent:
            url = "http://localhost:11434/v1/chat/completions"
            headers = {
                "Content-Type": "application/json"
            }
            data = {
            "model": "mistral",
            "messages": [{
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": f"Fully summarize the text delimited by triple backticks.Text:```{processedContent}```"
            }]}
            response = requests.post(url, headers=headers, json=data)
            absText = response.json()['choices'][0]['message']['content']
            # data = {
            # "model": "mistral",
            # "messages": [{
            #     "role": "system",
            #     "content": "You're a professional translator,translate the text delimited by triple backticks to Chinese."
            # },
            # {
            #     "role": "user",
            #     "content": f"```{absText}```"
            # }]}
            # response = requests.post(url, headers=headers, json=data)
            # absText_CHN = response.json()['choices'][0]['message']['content']
            ws.append([foundFile, absText])
        
        wb.save('output.xlsx')